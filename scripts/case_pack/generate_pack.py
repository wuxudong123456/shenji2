from pathlib import Path
import json, hashlib, csv
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4

ROOT=Path('testdata/government_procurement_full_case')
FONT=Path(r'C:\Windows\Fonts\msyh.ttc')
pdfmetrics.registerFont(TTFont('CN',str(FONT)))
facts={
 'project':{'name':'东河县教育局2025年度政府采购专项审计','unit':'东河县教育局（虚构）','period':'2025-01-01至2025-12-31'},
 'contracts':[
 ['A','DHJY-CG-2025-001','东河数智设备有限公司',1680000,'公开招标','2025-03-18','教学终端采购'],
 ['B1','DHJY-CG-2025-021','东河优居家具有限公司',620000,'询价','2025-05-08','办公家具采购（一）'],
 ['B2','DHJY-CG-2025-022','东河优居商贸有限公司',590000,'询价','2025-05-15','办公家具采购（二）'],
 ['B3','DHJY-CG-2025-023','东河优居供应链有限公司',610000,'询价','2025-05-22','办公家具采购（三）'],
 ['C','DHJY-CG-2025-031','东河云维科技有限公司',800000,'竞争性磋商','2025-07-01','信息系统运维服务']],
 'findings':[
 ['GP-001','拆分采购疑点','B1/B2/B3同属信息化提升预算项目，14日内同类家具合同合计182万元','高'],
 ['GP-002','超比例收取履约保证金','C合同80万元，履约保证金10万元，占比12.5%','高'],
 ['GP-003','验收时点异常','C项目验收日期2025-11-20，早于服务完成日2025-12-31','中'],
 ['GP-004','票款金额不一致','C项目发票80万元，实际付款78万元，差额2万元','中']],
}
dirs=['00_README','01_项目立项','02_审计输入','03_原始资料/A_正常采购','03_原始资料/B_拆分采购','03_原始资料/C_保证金及履约异常','04_结构化基准/csv','05_违规规则','06_标准答案','07_预期成果','08_操作手册','09_验收记录']
for d in dirs:(ROOT/d).mkdir(parents=True,exist_ok=True)

def docx(path,title,sections):
 d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=Cm(2.2); sec.left_margin=sec.right_margin=Cm(2.4)
 st=d.styles['Normal']; st.font.name='Microsoft YaHei'; st.font.size=Pt(10.5)
 p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True;r.font.size=Pt(18)
 p=d.add_paragraph('虚拟测试资料｜仅用于审计工坊功能验证');p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 for h,rows in sections:
  d.add_heading(h,level=1)
  if isinstance(rows,str): d.add_paragraph(rows)
  else:
   t=d.add_table(rows=0,cols=2);t.style='Table Grid'
   for a,b in rows:
    c=t.add_row().cells;c[0].text=str(a);c[1].text=str(b)
 d.save(path)

styles=getSampleStyleSheet(); styles.add(ParagraphStyle(name='CNBody',fontName='CN',fontSize=10,leading=16));styles.add(ParagraphStyle(name='CNTitle',fontName='CN',fontSize=18,leading=24,alignment=1))
def pdf(path,title,sections):
 story=[Paragraph(title,styles['CNTitle']),Paragraph('虚拟测试资料｜仅用于审计工坊功能验证',styles['CNBody']),Spacer(1,12)]
 for h,rows in sections:
  story += [Paragraph(h,styles['CNBody']),Spacer(1,4)]
  if isinstance(rows,str):story.append(Paragraph(rows,styles['CNBody']))
  else:
   data=[[Paragraph(str(a),styles['CNBody']),Paragraph(str(b),styles['CNBody'])] for a,b in rows]
   t=Table(data,colWidths=[120,350]);t.setStyle(TableStyle([('FONT',(0,0),(-1,-1),'CN'),('GRID',(0,0),(-1,-1),.5,colors.grey),('VALIGN',(0,0),(-1,-1),'TOP'),('BACKGROUND',(0,0),(0,-1),colors.HexColor('#E8EEF5'))]));story.append(t)
  story.append(Spacer(1,10))
 SimpleDocTemplate(str(path),pagesize=A4,rightMargin=45,leftMargin=45,topMargin=45,bottomMargin=45).build(story)

intro=[('项目基本信息',[('项目名称',facts['project']['name']),('被审计单位',facts['project']['unit']),('审计期间',facts['project']['period']),('审计目标','检查拆分采购、保证金、合同履行、验收、发票与付款合规性')])]
docx(ROOT/'01_项目立项/项目立项书.docx','政府采购专项审计项目立项书',intro);pdf(ROOT/'01_项目立项/项目立项书.pdf','政府采购专项审计项目立项书',intro)
for row in facts['contracts']:
 g,no,supplier,amt,method,date,item=row; folder={'A':'A_正常采购','B':'B_拆分采购','C':'C_保证金及履约异常'}[g[0]]
 sections=[('合同信息',[('合同编号',no),('采购人','东河县教育局（虚构）'),('供应商',supplier),('合同名称',item),('采购方式',method),('合同金额',f'{amt:,.2f}元'),('签订日期',date)]),('履约条款',[('履约保证金', '100000元（合同C）' if g=='C' else '不收取'),('服务/交付完成日','2025-12-31' if g=='C' else '按合同约定')])]
 docx(ROOT/f'03_原始资料/{folder}/{no}_采购合同.docx',f'{item}合同',sections);pdf(ROOT/f'03_原始资料/{folder}/{no}_采购合同.pdf',f'{item}合同',sections)

intent='对东河县教育局2025年度政府采购活动开展专项审计，重点检查是否通过拆分采购规避公开招标，是否超比例收取履约保证金，以及采购合同、验收、发票、记账凭证和银行付款之间是否一致。'
(ROOT/'02_审计输入/审计意图.txt').write_text(intent,encoding='utf-8')
rules='''# 违规规则\n\n- GP-001：采购汇总台账.同预算项目合同合计金额 > 1000000 AND 采购汇总台账.采购方式 = "询价"\n- GP-002：保证金台账.实际收取金额 / 保证金台账.合同金额 > 0.10\n- GP-003：验收台账.验收日期 < 验收台账.合同完成日期\n- GP-004：支付台账.发票金额 != 支付台账.实际付款金额\n\nGP-001 使用预计算汇总字段，适配当前逐行表达式引擎。'''
(ROOT/'05_违规规则/违规表达式与说明.md').write_text(rules,encoding='utf-8')
ans='# 标准答案\n\n预期命中4个核心疑点。A组为阴性样本，不应命中。\n\n'+'\n'.join(f'- {x[0]} {x[1]}：{x[2]}（{x[3]}）' for x in facts['findings'])
(ROOT/'06_标准答案/预期违规命中.md').write_text(ans,encoding='utf-8')
report=[('审计范围',[('单位',facts['project']['unit']),('期间',facts['project']['period'])]),('系统识别疑点',[(x[0]+' '+x[1],x[2]) for x in facts['findings']]),('结论边界','以上均为系统疑点，须由审计人员取得补充证据并履行复核程序后定性。')]
docx(ROOT/'07_预期成果/审计疑点报告.docx','政府采购专项审计疑点报告',report);pdf(ROOT/'07_预期成果/审计疑点报告.pdf','政府采购专项审计疑点报告',report)
guide=[('操作流程','1. 创建项目；2. 粘贴审计意图；3. 确认违规与法规；4. 上传PDF/DOCX/XLSX；5. 检查OCR和模板；6. 对照金标准；7. 执行表达式；8. 生成疑点并核查溯源。'),('关键接口',[('创建项目','POST /api/audit/projects'),('上传文件','POST /api/audit/projects/{project_id}/upload，表单字段file'),('创建分析','POST /api/audit/analysis'),('人工确认','POST /api/audit/analysis/{task_id}/confirm'),('执行表达式','POST /api/audit/expression/execute'),('生成疑点','POST /api/audit/suspicion/generate')]),('预期结果',[('核心疑点','4个'),('阴性样本','A组不命中'),('法规检索关键词','政府采购、公开招标、履约保证金、合同验收')])]
docx(ROOT/'08_操作手册/审计工坊完整业务流程操作手册.docx','审计工坊完整业务流程操作手册',guide);pdf(ROOT/'08_操作手册/审计工坊完整业务流程操作手册.pdf','审计工坊完整业务流程操作手册',guide)

contracts=[]
for g,no,sup,amt,meth,date,item in facts['contracts']:
 contracts.append({'project_id':'CASE-GP-2025','document_trace_id':f'TRACE-{g}-CONTRACT','template_name':'audit/合同协议类/合同','doc_name':no+'_采购合同.pdf','doc_type':'采购合同','party_a':'东河县教育局（虚构）','party_b':sup,'amount':amt,'currency':'CNY','sign_date':date,'contract_no':no,'procurement_method':meth,'extra_fields':json.dumps({'采购项目':item},ensure_ascii=False)})
gold={'project':facts['project'],'data_contracts':contracts,'findings':facts['findings']}
(ROOT/'04_结构化基准/golden_dataset.json').write_text(json.dumps(gold,ensure_ascii=False,indent=2),encoding='utf-8')
with open(ROOT/'04_结构化基准/csv/data_contracts.csv','w',newline='',encoding='utf-8-sig') as f:
 w=csv.DictWriter(f,fieldnames=contracts[0].keys());w.writeheader();w.writerows(contracts)
(ROOT/'00_README/资料包说明.md').write_text('# 审计工坊完整虚拟案例包\n\n本资料包用于政府采购全流程测试。所有单位、人员、编号和业务均为虚构。优先上传 `03_原始资料` 中的 PDF；用 `04_结构化基准` 校验提取结果，用 `06_标准答案` 核验规则命中。\n',encoding='utf-8')
(Path('scripts/case_pack/case_facts.json')).write_text(json.dumps(facts,ensure_ascii=False,indent=2),encoding='utf-8')
manifest=[]
for p in sorted(ROOT.rglob('*')):
 if p.is_file() and p.name!='manifest.json':manifest.append({'path':p.relative_to(ROOT).as_posix(),'size':p.stat().st_size,'sha256':hashlib.sha256(p.read_bytes()).hexdigest()})
(ROOT/'manifest.json').write_text(json.dumps({'case':'政府采购综合审计','file_count':len(manifest),'files':manifest},ensure_ascii=False,indent=2),encoding='utf-8')
print('PACK_OK',len(manifest))
