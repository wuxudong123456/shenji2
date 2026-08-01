from pathlib import Path
import json,csv,hashlib
from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer,Table,TableStyle
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

R=Path('testdata/real_source_adapted_procurement_case')
for d in ['00_使用说明','01_公开来源证据','02_真实事实与补建边界','03_审计意图','04_补建原始资料','05_结构化金标准/csv','06_违规规则','07_标准答案','08_疑点成果','09_数据库导入','10_操作手册']:(R/d).mkdir(parents=True,exist_ok=True)
SOURCE='https://www.audit.gov.cn/n5/n1482/c140763/part/73697.pdf'
LAW='https://gks.mof.gov.cn/guizhangzhidu/200807/t20080731_59933.htm'
REG='https://www.mof.gov.cn/zhengwuxinxi/zhengcefabu/201502/t20150227_1195516.htm'
rows=[
('U01','信息化设备采购',1860000,'直接采购','2025-02-18'),('U02','办公设备采购',1280000,'直接采购','2025-03-22'),
('U03','物业服务采购',1740000,'直接采购','2025-04-15'),('U04','实验设备采购',2150000,'直接采购','2025-06-05'),
('U05','印刷服务采购',960000,'直接采购','2025-07-11'),('U06','软件运维服务采购',1530000,'直接采购','2025-09-08'),
('U07','办公家具采购',1401400,'直接采购','2025-11-19')]
assert sum(x[2] for x in rows)==10921400

styles=getSampleStyleSheet();pdfmetrics.registerFont(TTFont('CN',r'C:\Windows\Fonts\msyh.ttc'));styles.add(ParagraphStyle(name='B',fontName='CN',fontSize=10,leading=16));styles.add(ParagraphStyle(name='T',fontName='CN',fontSize=18,leading=24,alignment=1))
def sections_for(unit,item,amt,date):return [('补建业务信息',[('匿名单位',unit),('采购事项',item),('采购金额',f'{amt:,.2f}元'),('采购方式','直接采购'),('发生日期',date),('政府采购程序','未执行')]),('材料属性',[('属性','根据公开审计汇总事实补建的测试材料'),('真实性说明','本文件不是原审计项目底稿，不代表公开公告披露了该笔明细'),('关联真实事实','7个单位应实行未实行政府采购，公开金额合计1092.14万元')])]
def docx(path,title,secs):
 d=Document();s=d.sections[0];s.top_margin=s.bottom_margin=Cm(2.2);s.left_margin=s.right_margin=Cm(2.4);d.styles['Normal'].font.name='Microsoft YaHei';d.styles['Normal'].font.size=Pt(10.5)
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(title);r.bold=True;r.font.size=Pt(18)
 p=d.add_paragraph('真实来源改编｜补建测试材料｜不得视为原始审计证据');p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 for h,data in secs:
  d.add_heading(h,1)
  if isinstance(data,str):d.add_paragraph(data)
  else:
   t=d.add_table(rows=0,cols=2);t.style='Table Grid'
   for a,b in data:c=t.add_row().cells;c[0].text=str(a);c[1].text=str(b)
 d.save(path)
def pdf(path,title,secs):
 story=[Paragraph(title,styles['T']),Paragraph('真实来源改编｜补建测试材料｜不得视为原始审计证据',styles['B']),Spacer(1,12)]
 for h,data in secs:
  story.append(Paragraph(h,styles['B']))
  if isinstance(data,str):story.append(Paragraph(data,styles['B']))
  else:
   t=Table([[Paragraph(str(a),styles['B']),Paragraph(str(b),styles['B'])] for a,b in data],colWidths=[120,350]);t.setStyle(TableStyle([('FONT',(0,0),(-1,-1),'CN'),('GRID',(0,0),(-1,-1),.5,colors.grey),('BACKGROUND',(0,0),(0,-1),colors.HexColor('#E8EEF5')),('VALIGN',(0,0),(-1,-1),'TOP')]));story.append(t)
  story.append(Spacer(1,10))
 SimpleDocTemplate(str(path),pagesize=A4,rightMargin=45,leftMargin=45,topMargin=45,bottomMargin=45).build(story)

source_md=f'''# 公开来源证据台账\n\n## 真实审计事实\n\n河南省2019年度省级预算执行和其他财政收支审计公告第4页披露：5个省直部门和2家所属单位采购商品或服务，应实行未实行政府采购，涉及采购金额1092.14万元。\n\n- 原始公告：{SOURCE}\n- 定位：PDF第4页，公开文本第90-92行附近\n- 发布主体：河南省审计机关（审计署网站转载）\n\n## 法规来源\n\n- 《中华人民共和国政府采购法》：{LAW}\n- 《中华人民共和国政府采购法实施条例》：{REG}\n\n## 使用限制\n\n公告未公开单位名称、单笔合同、供应商、发票、凭证和付款流水。资料包内相关内容均为测试补建，不得称为公告原始事实。\n'''
(R/'01_公开来源证据/公开来源证据台账.md').write_text(source_md,encoding='utf-8')
(R/'02_真实事实与补建边界/事实边界清单.md').write_text('# 事实边界\n\n| 内容 | 属性 |\n|---|---|\n| 5个省直部门、2家所属单位 | 公开真实事实 |\n| 应实行未实行政府采购 | 公开真实事实 |\n| 合计1092.14万元 | 公开真实事实 |\n| U01-U07、采购事项、日期和单笔金额 | 测试补建 |\n| 合同、发票、凭证、付款回单 | 测试补建 |\n',encoding='utf-8')
for unit,item,amt,method,date in rows:
 secs=sections_for(unit,item,amt,date);docx(R/f'04_补建原始资料/{unit}_{item}.docx',f'{unit} {item}资料摘要',secs);pdf(R/f'04_补建原始资料/{unit}_{item}.pdf',f'{unit} {item}资料摘要',secs)
intent='核查匿名单位U01至U07采购商品和服务是否依法履行政府采购程序，重点识别应实行未实行政府采购、直接采购以及可能规避公开招标的疑点，并核对合计金额是否为1092.14万元。'
(R/'03_审计意图/可粘贴审计意图.txt').write_text(intent,encoding='utf-8')
gold=[{'project_id':'REAL-ADAPT-2019-HN','document_trace_id':f'TRACE-{u}','template_name':'audit/合同协议类/合同','doc_name':f'{u}_{i}.pdf','doc_type':'采购资料摘要','party_a':u,'party_b':'匿名供应商（补建）','amount':a,'currency':'CNY','sign_date':d,'contract_no':f'ADAPT-{u}-2019','procurement_method':m,'extra_fields':json.dumps({'是否执行政府采购':'否','采购事项':i,'数据属性':'测试补建'},ensure_ascii=False)} for u,i,a,m,d in rows]
(R/'05_结构化金标准/golden_dataset.json').write_text(json.dumps({'source_fact':{'unit_count':7,'amount':10921400,'source':SOURCE},'data_contracts':gold},ensure_ascii=False,indent=2),encoding='utf-8')
with open(R/'05_结构化金标准/csv/data_contracts.csv','w',newline='',encoding='utf-8-sig') as f:w=csv.DictWriter(f,fieldnames=gold[0].keys());w.writeheader();w.writerows(gold)
(R/'06_违规规则/规则与表达式.md').write_text("# 规则\n\n- RA-001：合同数据.采购方式 = '直接采购' AND 合同数据.金额 > 0\n- RA-002：汇总台账.单位数量 = 7 AND 汇总台账.采购总金额 = 10921400\n\nRA-001预期命中7行；RA-002用于校验公开汇总事实。是否依法必须公开招标仍需结合当年度目录、限额标准和批准文件人工定性。\n",encoding='utf-8')
(R/'07_标准答案/预期命中.md').write_text('# 标准答案\n\nRA-001预期扫描7行、命中7行，命中金额合计10,921,400元。真实公开结论仅为“应实行未实行政府采购”；不得仅凭补建数据认定化整为零。\n',encoding='utf-8')
report=[('真实来源事实',[('涉及单位','5个省直部门和2家所属单位'),('问题','采购商品或服务应实行未实行政府采购'),('涉及金额','1092.14万元'),('来源',SOURCE)]),('系统测试结果',[('补建记录','7条'),('规则命中','7条'),('命中金额','1092.14万元')]),('审计结论边界','系统命中属于测试疑点。单位身份、单笔事实和责任认定须以原审计底稿或正式调查材料为准。')]
docx(R/'08_疑点成果/真实来源改编案例疑点报告.docx','真实来源改编案例疑点报告',report);pdf(R/'08_疑点成果/真实来源改编案例疑点报告.pdf','真实来源改编案例疑点报告',report)
sql="""-- 真实来源改编案例；明细为测试补建\nINSERT INTO audit_cases (title,domain,case_summary,audit_method,involved_amount,source,status) VALUES ('河南省2019年度应实行未实行政府采购问题（改编测试）','政府采购','公开公告披露5个省直部门和2家所属单位应实行未实行政府采购，涉及1092.14万元；单位及明细未公开。','公告事实提取+补建数据规则验证',10921400,'"+SOURCE+"','published');\n"""
(R/'09_数据库导入/audit_cases_import.sql').write_text(sql,encoding='utf-8')
guide=[('使用顺序','1. 阅读事实边界；2. 创建项目；3. 粘贴审计意图；4. 上传04目录PDF；5. 对照金标准；6. 执行RA-001；7. 核对7条、1092.14万元；8. 查看疑点报告。'),('系统接口',[('创建项目','POST /api/audit/projects'),('上传','POST /api/audit/projects/{id}/upload，字段file'),('表达式','POST /api/audit/expression/execute'),('疑点','POST /api/audit/suspicion/generate')]),('验收重点',[('OCR','能识别匿名单位、事项、金额、日期和补建标识'),('提取','模板建议audit/合同协议类/合同'),('规则','7/7命中，金额1092.14万元'),('合规','输出不得把补建明细称为真实原始事实')])]
docx(R/'10_操作手册/真实来源改编案例操作手册.docx','真实来源改编案例操作手册',guide);pdf(R/'10_操作手册/真实来源改编案例操作手册.pdf','真实来源改编案例操作手册',guide)
(R/'00_使用说明/README.md').write_text('# 真实来源改编审计案例包\n\n真实核心事实来自河南省审计公告；未公开的单位和交易附件均为补建测试数据。使用前必须先阅读“事实边界清单”。\n',encoding='utf-8')
manifest=[]
for p in sorted(R.rglob('*')):
 if p.is_file() and p.name!='manifest.json':manifest.append({'path':p.relative_to(R).as_posix(),'size':p.stat().st_size,'sha256':hashlib.sha256(p.read_bytes()).hexdigest()})
(R/'manifest.json').write_text(json.dumps({'name':'真实来源改编政府采购审计案例包','source':SOURCE,'file_count':len(manifest),'files':manifest},ensure_ascii=False,indent=2),encoding='utf-8')
print('REAL_PACK_OK',len(manifest),sum(x[2] for x in rows))
