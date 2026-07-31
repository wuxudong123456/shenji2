"""审计文书 Word(.docx) 导出服务

把 document_service.batch_generate 产出的结构化文书对象渲染成 .docx：
  1. 取证单 (evidence)
  2. 审计底稿 (workpaper)
  3. 审计报告 (report)
  4. 审理复核意见书 (review)

docx 写法参考 scripts/case_pack/generate_pack.py 的现成范式
（页边距 / 雅黑字体 / 居中标题 / Table Grid 表格）。
"""
import io
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from services.document_service import batch_generate, generate_document

_FONT = "Microsoft YaHei"
_TITLE_CN = {
    "evidence": "取证单",
    "workpaper": "审计底稿",
    "report": "审计报告",
    "review": "审理复核意见书",
}


# ──────────────────────────────────────────────
#  docx 公共工具
# ──────────────────────────────────────────────

def _new_doc() -> Document:
    """新建 Document，预设 Normal 样式（雅黑 / 10.5pt）与页边距。"""
    d = Document()
    sec = d.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.4)
    st = d.styles["Normal"]
    st.font.name = _FONT
    st.font.size = Pt(10.5)
    # 显式指定东亚字体，确保中文不走默认宋体回退
    st.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    return d


def _add_title(d: Document, text: str):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text or "")
    r.bold = True
    r.font.size = Pt(18)


def _add_section_heading(d: Document, text: str):
    p = d.add_paragraph()
    r = p.add_run(text or "")
    r.bold = True
    r.font.size = Pt(13)


def _add_meta_table(d: Document, rows):
    """两列信息表（编号/项目/日期…）。"""
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = str(k or "")
        cells[1].text = str(v if v is not None else "")


def _add_signature(d: Document):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("审计员：________　复核人：________　日期：________")


def _doc_to_bytesio(d: Document) -> io.BytesIO:
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
#  四类文书的 docx 渲染
# ──────────────────────────────────────────────

def _render_evidence(doc: dict) -> io.BytesIO:
    d = _new_doc()
    _add_title(d, doc.get("title", "审计取证单"))
    _add_meta_table(d, [
        ("编号", doc.get("code", "")),
        ("项目", doc.get("project", "")),
        ("审计员", doc.get("auditor", "")),
        ("日期", doc.get("date", "")),
    ])
    _add_section_heading(d, "一、取证事项")
    items = doc.get("audit_items") or []
    if not items:
        d.add_paragraph("（暂无取证事项，请在前面步骤确认疑点与法规）")
    for i, it in enumerate(items, 1):
        d.add_paragraph(f"{i}. {it.get('audit_item', '')}")
        finding = it.get("finding", "") or ""
        if it.get("amount"):
            finding += f"（涉及金额：{it.get('amount')}）"
        if it.get("period"):
            finding += f"（涉及期间：{it.get('period')}）"
        d.add_paragraph(finding)
        basis = "；".join(
            f"{l.get('law', '')}{('·' + l['clause']) if l.get('clause') else ''}"
            for l in (it.get("legal_basis") or [])
        )
        d.add_paragraph(f"法规依据：{basis or '—'}")
    _add_signature(d)
    return _doc_to_bytesio(d)


def _render_workpaper(doc: dict) -> io.BytesIO:
    d = _new_doc()
    _add_title(d, doc.get("title", "审计工作底稿"))
    _add_meta_table(d, [("编号", doc.get("code", "")), ("日期", doc.get("date", ""))])

    _add_section_heading(d, "一、审计程序")
    for i, p in enumerate(doc.get("procedures") or [], 1):
        d.add_paragraph(f"{i}. {p}")

    _add_section_heading(d, "二、审计发现")
    d.add_paragraph(doc.get("findings") or "—")

    _add_section_heading(d, "三、证据链")
    ev = doc.get("evidence_list") or []
    if ev:
        for i, e in enumerate(ev, 1):
            d.add_paragraph(f"{i}. {e}")
    else:
        d.add_paragraph("—")
    _add_signature(d)
    return _doc_to_bytesio(d)


def _render_report(doc: dict) -> io.BytesIO:
    d = _new_doc()
    _add_title(d, doc.get("title", "审计报告"))
    _add_meta_table(d, [("编号", doc.get("code", "")), ("日期", doc.get("date", ""))])

    _add_section_heading(d, "一、审计概况")
    d.add_paragraph(doc.get("summary") or "—")

    _add_section_heading(d, "二、疑点统计")
    d.add_paragraph(f"共 {doc.get('total_suspicions', 0)} 条，其中高风险 {doc.get('high_risk_count', 0)} 条。")

    _add_section_heading(d, "三、主要疑点")
    sus = doc.get("suspicions") or []
    if sus:
        for i, s in enumerate(sus, 1):
            title = s.get("violation_title") or s.get("title") or s.get("name") or f"疑点{i}"
            desc = s.get("description") or ""
            d.add_paragraph(f"{i}. {title}" + (f"：{desc}" if desc else ""))
    else:
        d.add_paragraph("—")

    _add_section_heading(d, "四、审计建议")
    for i, r in enumerate(doc.get("recommendations") or [], 1):
        d.add_paragraph(f"{i}. {r}")
    _add_signature(d)
    return _doc_to_bytesio(d)


def _render_review(doc: dict) -> io.BytesIO:
    d = _new_doc()
    _add_title(d, doc.get("title", "审理复核意见书"))
    _add_meta_table(d, [("编号", doc.get("code", "")), ("日期", doc.get("date", ""))])

    t = d.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "序号", "复核项", "AI 评估", "人工复核"
    for i, it in enumerate(doc.get("review_items") or [], 1):
        cells = t.add_row().cells
        cells[0].text = str(i)
        cells[1].text = it.get("item", "")
        cells[2].text = it.get("ai_assessment", "")
        cells[3].text = it.get("human_review", "") or "（待填写）"
    _add_signature(d)
    return _doc_to_bytesio(d)


_RENDERERS = {
    "evidence": _render_evidence,
    "workpaper": _render_workpaper,
    "report": _render_report,
    "review": _render_review,
}


# ──────────────────────────────────────────────
#  安全获取 doc 对象（report 的 LLM 失败降级）
# ──────────────────────────────────────────────

def _fallback_report(context: dict) -> dict:
    """report 在 LLM 不可用时的无推理占位（与 _build_report_template 失败分支同构）。"""
    suspicions = context.get("suspicions", [])
    return {
        "doc_type": "report",
        "title": f"审计报告 — {context.get('project_title', '')}",
        "code": f"AR-{datetime.now().strftime('%Y%m%d-%H%M')}",
        "summary": context.get("analysis_summary") or "（AI 推理暂不可用，已回退到分析摘要）",
        "suspicions": suspicions,
        "total_suspicions": len(suspicions),
        "high_risk_count": 0,
        "recommendations": [
            "建议被审计单位针对上述问题逐项整改",
            "完善内部控制制度，堵塞管理漏洞",
            "对涉及违规资金依法依规处理",
        ],
        "generated_at": datetime.now().isoformat(),
    }


def _safe_batch_generate(context: dict) -> dict:
    """复用 batch_generate；整体或 report 失败时逐项降级，保证导出永不因 LLM 而整批失败。"""
    try:
        result = batch_generate(context)
        if result.get("success") and result.get("documents"):
            docs = dict(result["documents"])
            if "report" not in docs:
                docs["report"] = _fallback_report(context)
            return docs
    except Exception:
        pass
    # 降级：逐项构建（非 report 无 LLM），report 走占位
    docs = {}
    for t in ("evidence", "workpaper", "review"):
        try:
            r = generate_document(t, context)
            if r.get("success"):
                docs[t] = r["document"]
        except Exception:
            continue
    docs["report"] = _fallback_report(context)
    return docs


# ──────────────────────────────────────────────
#  对外入口
# ──────────────────────────────────────────────

def export_single(doc_type: str, context: dict):
    """导出单文书 → (BytesIO, 文件名)。

    非报告类型只生成该类型（不触发 report 的 LLM）；报告类型走 _safe_batch_generate 降级。
    """
    if doc_type not in _RENDERERS:
        raise ValueError(f"不支持的文书类型: {doc_type}")

    doc = None
    if doc_type != "report":
        try:
            r = generate_document(doc_type, context)
            if r.get("success"):
                doc = r["document"]
        except Exception:
            doc = None
    if doc is None:
        doc = _safe_batch_generate(context).get(doc_type)
    if not doc:
        raise ValueError(f"文书 {doc_type} 生成失败，请检查上下文数据")

    buf = _RENDERERS[doc_type](doc)
    stamp = datetime.now().strftime("%Y%m%d")
    filename = f"{_TITLE_CN[doc_type]}-{stamp}.docx"
    return buf, filename


def export_all(context: dict):
    """导出四件套 → (zip BytesIO, 文件名)。"""
    docs = _safe_batch_generate(context)
    zip_buf = io.BytesIO()
    stamp = datetime.now().strftime("%Y%m%d")
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc_type in ("evidence", "workpaper", "report", "review"):
            doc = docs.get(doc_type)
            if not doc:
                continue
            d_buf = _RENDERERS[doc_type](doc)
            zf.writestr(f"{_TITLE_CN[doc_type]}-{stamp}.docx", d_buf.getvalue())
    zip_buf.seek(0)
    return zip_buf, f"审计文书四件套-{stamp}.zip"
